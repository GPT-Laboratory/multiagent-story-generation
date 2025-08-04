import asyncio
from concurrent.futures import ThreadPoolExecutor
import datetime
import io
import os
import random
import re
from bson import ObjectId
from dotenv import load_dotenv
from docx.shared import Pt

import asyncio
import io
import httpx
import random
from fastapi import FastAPI, Request, Response
from fastapi.responses import JSONResponse
from motor.motor_asyncio import AsyncIOMotorClient
import motor.motor_asyncio as motor_asyncio
# import motor_gridfs
from docx import Document
from docx.shared import Pt

import asyncio
import httpx
import io
import tempfile
import os
from docx import Document
from docx.shared import Pt

from fastapi import Request, Response
import openai
import base64
from fpdf import FPDF
import requests
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
load_dotenv()

from fastapi.responses import JSONResponse, Response
import io
from docx import Document
from motor.motor_asyncio import AsyncIOMotorClient
from motor.motor_asyncio import AsyncIOMotorGridFSBucket
from gridfs import GridFS
# Print all environment variables to verify
print("All environment variables:")
for key, value in os.environ.items():
    print(f"{key}: {value}")

OPENAI_API_KEY = os.getenv("API-KEY1")
LLAMA_API_KEY = os.getenv("LLAMA-key1")
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")
OPENAI_URL = "https://api.openai.com/v1/chat/completions"
LLAMA_URL="https://api.groq.com/openai/v1/chat/completions"
# DEEPSEEK_URL="https://gptlab.rd.tuni.fi/GPT-Lab/resources/Otula-P40-L4/api/generate",
DEEPSEEK_URL = "https://openrouter.ai/api/v1/chat/completions"

GPT_API_KEY = os.getenv("GPT_API_KEY")
print(f"DeepSeek API Key: {os.getenv('DEEPSEEK_API_KEY')}")
# Load environment variables from .env file
api_keys = [os.getenv(f"API-KEY{i}") for i in range(1, 4)]
llama_keys = [os.getenv(f"LLAMA-key{i}") for i in range(1, 3)]


MONGO_URI = os.getenv("MONGO_URI")
client = AsyncIOMotorClient(MONGO_URI)
db = client["MVP"]  # Database name
project_reports = db["project_reports"]
fs = AsyncIOMotorGridFSBucket(db)




def sanitize_text(text):
    """
    Replaces or removes unsupported Unicode characters in the text.
    If the input is a list, sanitize each item in the list.
    """
    # Define a dictionary of replacements for common Unicode characters
    replacements = {
        "\u2019": "'",  # Right single quotation mark
        "\u2018": "'",  # Left single quotation mark
        "\u201C": '"',  # Left double quotation mark
        "\u201D": '"',  # Right double quotation mark
        "\u2013": "-",  # En dash
        "\u2014": "--", # Em dash
        "\u2026": "...", # Ellipsis
        "\u00A0": " ",  # Non-breaking space
        # Add more replacements as needed
    }

    if isinstance(text, list):
        # If the input is a list, sanitize each item
        return [sanitize_text(item) for item in text]
    elif isinstance(text, str):
        # If the input is a string, sanitize it
        for char, replacement in replacements.items():
            text = text.replace(char, replacement)
        # Optionally, remove any remaining unsupported characters
        text = text.encode("latin-1", errors="ignore").decode("latin-1")
        return text
    else:
        # If the input is neither a string nor a list, return it as-is
        return text

# for document viewer 
async def get_report(request):
    file_id = request.path_params["file_id"]

    report = await project_reports.find_one({"file_id": file_id})
    if not report:
        return Response(content="Report not found", status_code=404)

    file_id = report["file_id"]
    fs = AsyncIOMotorGridFSBucket(db)
    file_stream = await fs.open_download_stream(ObjectId(file_id))

    file_data = await file_stream.read()

    return Response(
        content=file_data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": f"inline; filename={report['title'].replace(' ', '_')}.docx"
        }
    )


# fetch all reports
async def get_all_reports(request):
    selected_user_story_id = request.path_params["selectedUserStoryId"]

    reports_cursor = project_reports.find({"selected_user_story_id": selected_user_story_id})
    reports = await reports_cursor.to_list(length=None)

    if not reports:
        return JSONResponse({"message": "No reports found"}, status_code=404)

    # Convert ObjectId to string for JSON serialization
    for report in reports:
        report["_id"] = str(report["_id"])
        report["file_id"] = str(report["file_id"])

    return JSONResponse({"reports": reports})


def generate_ai_report_prompt(title, roles, mvp, vision, prioritized_stories, feedback, prioritizationResponses):
    """
    Generates a dynamic prompt for AI to generate the roadmap report based on user feedback.
    Only includes sections that are relevant according to the feedback provided.

    Args:
        title (str): Project title
        roles (list): List of agent roles involved
        mvp (str): Minimum Viable Product description
        vision (str): Project vision statement
        prioritized_stories (str): Prioritized user stories
        feedback (str): Feedback that determines which sections to include
        prioritizationResponses (dict): Agent prioritization responses

    Returns:
        str: Fully formatted prompt with only required sections based on feedback
    """
    # Start with the basic project information
    prompt_parts = [
        f"Create an AI Roadmap & Use Case Report for project '{title}' with the following details:",
        f"- Involved agents: {', '.join(roles)}"
    ]
    
    # Add sections based on what's likely needed from feedback
    if any(keyword in feedback.lower() for keyword in ['summary', 'overview', 'executive']):
        prompt_parts.append("""
##  Executive Summary
- Create a concise overview of the project
- Highlight key points from MVP and Vision
- Summarize the prioritized stories
""")
    
    if any(keyword in feedback.lower() for keyword in ['scope', 'definition', 'challenges']):
        prompt_parts.append(f"""
##  Scope Definition
- Analyze the vision: "{vision}"
- Examine the MVP: "{mvp}"
- Identify key challenges from: {prioritized_stories}
""")
    
    if any(keyword in feedback.lower() for keyword in ['ownership', 'requirement', 'assign']):
        prompt_parts.append(f"""
##  Requirement Ownership & Analysis
- Assign requirements to owners based on: {prioritizationResponses}
- Analyze any dependencies
""")
    
    
    if prioritizationResponses and any(keyword in feedback.lower() for keyword in ['agent', 'prioritization', 'comparison']):
        prompt_parts.append("""
## Prioritization Process Analysis
- Compare agent prioritization approaches
- Highlight agreements/disagreements
- Provide consensus recommendations
""")
    
    # Add the feedback instruction if provided
    if feedback:
        prompt_parts.insert(1, f"📌 **SPECIFIC FEEDBACK TO ADDRESS:**\n{feedback}\n")
    
    # Add formatting requirements
    prompt_parts.append("""
📌 **Formatting Requirements:**
- Use ONLY `##` for section headings
- Be concise and focused on the feedback
- Use bullet points (-) for lists
- Only include information derived from provided inputs
""")
    
    return "\n".join(prompt_parts).strip()


async def createDOCXReport(request: Request):
    data = await request.json()
    headers = {
        "Authorization": f"Bearer {random.choice(api_keys)}",
        "Content-Type": "application/json"
    }
    
    model = data['model']
    title = data.get("project_title", "AI Report")
    vision = data.get("vision", "No vision provided")
    mvp = data.get("mvp", "No MVP details provided")
    feedback = data.get("feedback", "")
    agents = data.get("agents", [])
    roles = [agent.get("role", "Unknown Role") for agent in agents]
    prioritized_stories = data.get("prioritized_stories", [])
    selected_user_story_id = data.get("selectedUserStory")
    project_id = data['project_id']
    prioritizationResponses = data.get("prioritizationResponses")
    
    print("prioritizationResponses:", prioritizationResponses)
    
    prompt = generate_ai_report_prompt(title, roles, mvp, vision, prioritized_stories, feedback, prioritizationResponses)
    
    if model in ["llama3-70b-8192", "mixtral-8x7b-32768", "deepseek-r1-distill-llama-70b-specdec"]:
        url = LLAMA_URL
        headers["Authorization"] = f"Bearer {LLAMA_API_KEY}"
    elif model == "deepseek/deepseek-r1-distill-llama-70b":
        url = DEEPSEEK_URL
        headers["Authorization"] = f"Bearer {DEEPSEEK_API_KEY}"
    else:
        url = OPENAI_URL
        headers["Authorization"] = f"Bearer {OPENAI_API_KEY}"
    
    post_data = {
        "model": model,
        "messages": [{"role": "user", "content": prompt},{"role": "user", "content": feedback}],
        "temperature": 0.7
    }
    
    # Use httpx for async HTTP requests
    async with httpx.AsyncClient() as client:
        try:
            response = await client.post(url, json=post_data, headers=headers, timeout=30)
            response.raise_for_status()
            response_data = response.json()
        except Exception as e:
            return Response(content=f"API request failed: {str(e)}", status_code=500)
    
    try:
        report_text = response_data["choices"][0]["message"]["content"]
    except KeyError:
        return Response(content="Invalid response format from the API", status_code=500)
    
    doc = Document()
    doc.add_heading(title, level=0)
   
    

    
    cleaned_report = report_text.replace("\\*\\*", "**")
    
    # Now process the cleaned report line by line
    for line in cleaned_report.split("\n"):
        paragraph = doc.add_paragraph()

        
        
        # Headers
        if line.startswith("# "):  # H1
            paragraph.add_run(line[2:].strip("*")).bold = True
            paragraph.runs[0].font.size = Pt(16)
        elif line.startswith("## "):  # H2
            paragraph.add_run(line[3:].strip("*")).bold = True
            paragraph.runs[0].font.size = Pt(14)
        # Bold text with asterisks
        elif "**" in line:
            # Process line with bold text
            parts = line.split("**")
            is_bold = False
            
            for i, part in enumerate(parts):
                if i % 2 == 0:  # Even parts (not bold)
                    if part:  # Don't add empty runs
                        paragraph.add_run(part)
                else:  # Odd parts (bold)
                    if part:  # Don't add empty runs
                        run = paragraph.add_run(part)
                        run.bold = True
        # Lists with bold prefixes
        elif line.startswith("- **") and "**:" in line:
            prefix_end = line.find("**:")
            prefix = line[3:prefix_end]
            content = line[prefix_end+3:]  # Skip the "**: "
            
            run = paragraph.add_run("- ")
            bold_run = paragraph.add_run(prefix + ": ")
            bold_run.bold = True
            paragraph.add_run(content)
        # Regular list items
        elif line.startswith("- "):
            paragraph.add_run(line)
        # Tables
        elif line.startswith("|"):
            paragraph.add_run(line)
        # Regular text
        else:
            paragraph.add_run(line)
     

    doc.add_heading("Kendall Tau Distance by Agent", level=1)

    if prioritizationResponses:
            # Dictionary to store best values for each agent type
            agent_metrics = {}
            
            # Process each response in the list
            for response in prioritizationResponses:
                if isinstance(response, dict) and 'message' in response and 'agentType' in response:
                    agent_type = response['agentType']
                    message = response['message']
                    
                    # Skip the "Final Prioritization" agent
                    if agent_type == "Final Prioritization":
                        continue
                        
                    # Initialize metrics for this agent if not already done
                    if agent_type not in agent_metrics:
                        agent_metrics[agent_type] = {
                            'best_avg_round': None,
                            'best_avg_value': float('inf'),
                            'best_pair_round': None,
                            'best_pair_value': float('inf')
                        }
                    
                    lines = message.split('\n')
                    
                    # Process Average Kendall Tau Distance table
                    in_avg_table = False
                    for line in lines:
                        if "| Round  | Average Kendall Tau Distance |" in line or "| Round | Average Kendall Tau Distance |" in line:
                            in_avg_table = True
                            continue
                        
                        if in_avg_table and "|" in line:
                            # Skip header separator
                            if "---" in line:
                                continue
                            
                            parts = line.split("|")
                            if len(parts) >= 3:
                                try:
                                    round_str = parts[1].strip()
                                    round_num = round_str.replace("Round ", "")
                                    value_str = parts[2].strip()
                                    value = float(value_str)
                                    
                                    # Update best value if this is lower
                                    if value < agent_metrics[agent_type]['best_avg_value']:
                                        agent_metrics[agent_type]['best_avg_value'] = value
                                        agent_metrics[agent_type]['best_avg_round'] = round_num
                                except (ValueError, IndexError):
                                    continue
                    
                    # Process Pairwise Kendall Tau Distance table
                    in_pair_table = False
                    for line in lines:
                        if "| Round Pair | Kendall Tau Distance |" in line:
                            in_pair_table = True
                            continue
                        
                        if in_pair_table and "|" in line:
                            # Skip header separator
                            if "---" in line:
                                continue
                            
                            parts = line.split("|")
                            if len(parts) >= 3:
                                try:
                                    pair_name = parts[1].strip()
                                    value_str = parts[2].strip()
                                    value = float(value_str)
                                    
                                    # Update best value if this is lower
                                    if value < agent_metrics[agent_type]['best_pair_value']:
                                        agent_metrics[agent_type]['best_pair_value'] = value
                                        agent_metrics[agent_type]['best_pair_round'] = pair_name
                                except (ValueError, IndexError):
                                    continue
            
                    # Create summary table
                    summary_table = doc.add_table(rows=1, cols=5)
                    summary_table.style = 'Table Grid'
                    
                    # Set header row
                    header_cells = summary_table.rows[0].cells
                    header_cells[0].text = "Agent Type"
                    header_cells[1].text = "Best Average Round"
                    header_cells[2].text = "Best Average Value"
                    header_cells[3].text = "Best Pair"
                    header_cells[4].text = "Best Pair Value"
                    
                    # Make header bold
                    for i in range(5):
                        for paragraph in header_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # Add data rows for each agent
                    for agent_type, metrics in agent_metrics.items():
                        row_cells = summary_table.add_row().cells
                        row_cells[0].text = agent_type
                        
                        # Best Average Round
                        if metrics['best_avg_round'] is not None:
                            row_cells[1].text = f"Round {metrics['best_avg_round']}"
                        else:
                            row_cells[1].text = "N/A"
                            
                        # Best Average Value
                        if metrics['best_avg_value'] != float('inf'):
                            row_cells[2].text = f"{metrics['best_avg_value']:.3f}"
                        else:
                            row_cells[2].text = "N/A"
                            
                        # Best Pair
                        if metrics['best_pair_round'] is not None:
                            row_cells[3].text = f"{metrics['best_pair_round']}"
                        else:
                            row_cells[3].text = "N/A"
                            
                        # Best Pair Value
                        if metrics['best_pair_value'] != float('inf'):
                            row_cells[4].text = f"{metrics['best_pair_value']:.3f}"
                        else:
                            row_cells[4].text = "N/A"
    else:
            doc.add_paragraph("No Kendall Tau Metrics available.")

            doc.add_heading("Agents Involved in Prioritization", level=1)
            doc.add_paragraph(', '.join(roles))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

   
    def upload_to_gridfs():
        # Use synchronous pymongo here instead of motor
        from pymongo import MongoClient
        import gridfs
        
        sync_client = MongoClient(MONGO_URI)
        # sync_db = sync_client.db
        sync_db = sync_client["MVP"]
        sync_fs = gridfs.GridFS(sync_db)
        
        file_id = sync_fs.put(
            buffer.getvalue(),
            filename=f"{title.replace(' ', '_')}.docx",
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        return file_id
    
    loop = asyncio.get_event_loop()
    with ThreadPoolExecutor() as executor:
        file_id = await loop.run_in_executor(executor, upload_to_gridfs)

    # Store DOCX in MongoDB GridFS
    file_id = await fs.upload_from_stream(
        f"{title.replace(' ', '_')}.docx",
        buffer.getvalue(),
        metadata={"content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    )

    # Store metadata in `project_reports` collection
    report_document = {
        "project_id": project_id,
        "selected_user_story_id": selected_user_story_id,
        "title": title,
        "file_id": str(file_id),  # Store file reference
    }

    await project_reports.insert_one(report_document)
    return JSONResponse(
        content={
            "message": "Report successfully created",
            "file_id": str(file_id),
            "title": title,
            "project_id": project_id
        },
        status_code=201
    )

async def createPDFReport(request: Request):
    data = await request.json()
    headers = {
        "Authorization": f"Bearer {random.choice(api_keys)}",
        "Content-Type": "application/json"
    }

    model = data['model']
    title = data.get("project_title", "AI Report")
    vision = data.get("vision", "No vision provided")
    mvp = data.get("mvp", "No MVP details provided")
    feedback = data.get("feedback", "")
    agents = data.get("agents", [])
    roles = [agent.get("role", "Unknown Role") for agent in agents]
    prioritized_stories = data.get("prioritized_stories", [])

    # Debugging Output
    print("Prioritized Stories:", prioritized_stories)

    # Generate AI Report Prompt
    prompt = generate_ai_report_prompt(title, roles, mvp, vision, prioritized_stories)

    # Choose API URL based on the model
    api_urls = {
        "llama3-70b-8192": LLAMA_URL,
        "mixtral-8x7b-32768": LLAMA_URL,
        "deepseek-r1-distill-llama-70b-specdec": LLAMA_URL,
        "deepseek/deepseek-r1-distill-llama-70b": DEEPSEEK_URL
    }
    url = api_urls.get(model, OPENAI_URL)
    headers["Authorization"] = f"Bearer {OPENAI_API_KEY if url == OPENAI_URL else LLAMA_API_KEY}"

    post_data = {
        "model": model,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.7
    }

    try:
        response = requests.post(url, json=post_data, headers=headers, timeout=30)
        response.raise_for_status()
        response_data = response.json()
    except requests.exceptions.RequestException as e:
        return Response(content=f"API request failed: {str(e)}", status_code=500)

    try:
        report_text = response_data["choices"][0]["message"]["content"]
    except KeyError:
        return Response(content="Invalid response format from the API", status_code=500)

    # Create PDF Document in Memory
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)
    width, height = letter
    y_position = height - 50

    def add_heading(text, font_size=14, y_offset=30):
        """Helper function to add a heading"""
        nonlocal y_position
        pdf.setFont("Helvetica-Bold", font_size)
        pdf.drawString(50, y_position, text)
        y_position -= y_offset

    def add_paragraph(text, font_size=12, line_spacing=16, max_width=500):
        """Helper function to add wrapped paragraphs"""
        nonlocal y_position
        pdf.setFont("Helvetica", font_size)
        lines = simpleSplit(text, "Helvetica", font_size, max_width)
        for line in lines:
            pdf.drawString(50, y_position, line)
            y_position -= line_spacing
            if y_position < 50:  # Handle page breaks
                pdf.showPage()
                pdf.setFont("Helvetica", font_size)
                y_position = height - 50

    # Add Title
    pdf.setFont("Helvetica-Bold", 18)
    pdf.drawString(50, y_position, title)
    y_position -= 40

    # Add MVP Section
    add_heading("1️⃣ MVP")
    add_paragraph(mvp)

    # Add Prioritized Stories
    if prioritized_stories:
        add_heading("2️⃣ Prioritized Stories")
        for story in prioritized_stories:
            add_paragraph(f"• {story.get('description', 'No description available')}")

    # Add Generated Report
    add_heading("3️⃣ Generated Report")
    for line in report_text.split("\n"):
        add_paragraph(line)

    # Finalize PDF
    pdf.save()
    buffer.seek(0)

    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename={title.replace(' ', '_')}.pdf"}
    )
